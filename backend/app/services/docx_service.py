"""Geração de DOCX para planejamentos de marketing."""

import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_planejamento_docx(
    cliente: dict,
    planejamento: dict,
    conteudos: list[dict],
) -> bytes:
    """Generate a DOCX file for a planning document."""
    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Header
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run(f"PLANEJAMENTO DE CONTEÚDO — {planejamento.get('mes_referencia', '').upper()}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(cliente.get("nome_empresa", ""))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

    doc.add_paragraph("")  # spacer

    # Resumo Estratégico
    if planejamento.get("resumo_estrategico"):
        _add_section_title(doc, "RESUMO ESTRATÉGICO")
        p = doc.add_paragraph(planejamento["resumo_estrategico"])
        p.paragraph_format.left_indent = Cm(0.5)

    # Temas
    temas = planejamento.get("temas") or []
    if temas:
        _add_section_title(doc, "TEMAS DO MÊS")
        for tema in temas:
            if isinstance(tema, dict):
                p = doc.add_paragraph()
                run = p.add_run(f"• {tema.get('tema', '')}")
                run.bold = True
                if tema.get("pilar"):
                    run = p.add_run(f"  [{tema['pilar']}]")
                    run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
                    run.font.size = Pt(9)
                if tema.get("justificativa"):
                    p2 = doc.add_paragraph(f"  {tema['justificativa']}")
                    p2.paragraph_format.left_indent = Cm(1)
                    p2.style.font.size = Pt(10)

    # Conteúdos por tipo
    videos = [c for c in conteudos if c.get("tipo") == "video_roteiro"]
    artes = [c for c in conteudos if c.get("tipo") == "arte_estatica"]
    carrosseis = [c for c in conteudos if c.get("tipo") == "carrossel"]

    # Vídeos
    if videos:
        _add_section_title(doc, "ROTEIROS DE VÍDEO")
        for i, video in enumerate(videos, 1):
            cont = video.get("conteudo", {})
            _add_content_header(doc, f"VÍDEO {i}", video.get("titulo", ""), video.get("framework"), video.get("pilar"))

            if cont.get("gancho"):
                p = doc.add_paragraph()
                run = p.add_run("GANCHO (0-3s): ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
                run.font.size = Pt(10)
                p.add_run(cont["gancho"])

            if cont.get("desenvolvimento"):
                p = doc.add_paragraph()
                run = p.add_run("DESENVOLVIMENTO: ")
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(cont["desenvolvimento"])

            if cont.get("cta") or cont.get("cta_final"):
                p = doc.add_paragraph()
                run = p.add_run("CTA: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
                run.font.size = Pt(10)
                p.add_run(cont.get("cta") or cont.get("cta_final", ""))

            if cont.get("duracao_estimada"):
                p = doc.add_paragraph(f"Duração estimada: {cont['duracao_estimada']}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            _add_variacoes(doc, video.get("variacoes_ab"))
            doc.add_paragraph("")

    # Artes
    if artes:
        _add_section_title(doc, "ARTES ESTÁTICAS")
        for i, arte in enumerate(artes, 1):
            cont = arte.get("conteudo", {})
            _add_content_header(doc, f"ARTE {i}", arte.get("titulo", ""), arte.get("framework"), arte.get("pilar"))

            if cont.get("titulo_arte") or cont.get("titulo"):
                p = doc.add_paragraph()
                run = p.add_run(cont.get("titulo_arte") or cont.get("titulo", ""))
                run.bold = True
                run.font.size = Pt(13)

            if cont.get("copy"):
                doc.add_paragraph(cont["copy"])

            if cont.get("cta_botao") or cont.get("cta"):
                p = doc.add_paragraph()
                run = p.add_run(f"CTA botão: {cont.get('cta_botao') or cont.get('cta', '')}")
                run.bold = True
                run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

            _add_variacoes(doc, arte.get("variacoes_ab"))
            doc.add_paragraph("")

    # Carrosséis
    if carrosseis:
        _add_section_title(doc, "CARROSSÉIS")
        for i, carrossel in enumerate(carrosseis, 1):
            cont = carrossel.get("conteudo", {})
            _add_content_header(doc, f"CARROSSEL {i}", carrossel.get("titulo", ""), carrossel.get("framework"), carrossel.get("pilar"))

            # Capa
            if cont.get("capa"):
                capa = cont["capa"]
                p = doc.add_paragraph()
                run = p.add_run("CAPA: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
                run.font.size = Pt(10)
                if isinstance(capa, dict):
                    p.add_run(capa.get("titulo", ""))
                    if capa.get("subtitulo"):
                        p.add_run(f" — {capa['subtitulo']}")
                else:
                    p.add_run(str(capa))

            # Slides
            slides = cont.get("slides", [])
            if isinstance(slides, list):
                for j, slide in enumerate(slides, 1):
                    p = doc.add_paragraph()
                    run = p.add_run(f"Slide {j}: ")
                    run.bold = True
                    run.font.size = Pt(10)
                    if isinstance(slide, dict):
                        if slide.get("titulo"):
                            p.add_run(slide["titulo"])
                        texto = slide.get("texto") or slide.get("conteudo", "")
                        if texto:
                            p2 = doc.add_paragraph(texto)
                            p2.paragraph_format.left_indent = Cm(1)
                    else:
                        p.add_run(str(slide))

            # CTA final
            if cont.get("cta_final"):
                p = doc.add_paragraph()
                run = p.add_run("CTA FINAL: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
                p.add_run(cont["cta_final"])

            # Legenda
            if cont.get("copy_legenda"):
                p = doc.add_paragraph()
                run = p.add_run("Legenda: ")
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(cont["copy_legenda"])

            _add_variacoes(doc, carrossel.get("variacoes_ab"))
            doc.add_paragraph("")

    # Calendário
    calendario = planejamento.get("calendario") or []
    if calendario:
        _add_section_title(doc, "CALENDÁRIO DE PUBLICAÇÃO")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Data"
        hdr[1].text = "Tipo"
        hdr[2].text = "Título"
        for item in calendario:
            if isinstance(item, dict):
                row = table.add_row().cells
                row[0].text = item.get("data", "")
                row[1].text = (item.get("tipo_conteudo") or item.get("tipo", "")).replace("_", " ")
                row[2].text = item.get("titulo", "")

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PMAX Agência Digital | Planejamento gerado com IA")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_section_title(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(18)
    # Add accent line (via border-bottom)
    from docx.oxml.ns import qn
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "12",
        qn("w:space"): "1",
        qn("w:color"): "E94560",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_content_header(doc: Document, number: str, title: str, framework: str | None, pilar: str | None):
    p = doc.add_paragraph()
    run = p.add_run(f"{number} — ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
    run.font.size = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    if framework:
        run = p.add_run(f"  [{framework}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    if pilar:
        run = p.add_run(f"  • {pilar}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_variacoes(doc: Document, variacoes: list | None):
    if not variacoes:
        return
    p = doc.add_paragraph()
    run = p.add_run("Variações A/B:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    for v in variacoes:
        if isinstance(v, dict):
            texto = v.get("copy_alternativa", str(v))
        else:
            texto = str(v)
        p2 = doc.add_paragraph(f"  → {texto}")
        p2.paragraph_format.left_indent = Cm(1)
        p2.runs[0].font.size = Pt(10)
